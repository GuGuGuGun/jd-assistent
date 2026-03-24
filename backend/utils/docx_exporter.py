"""
Word 文档导出模块，用于将 RenderReadyResume 模型渲染为 .docx 文件供下载
"""

import io

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .text_sanitizer import sanitize_resume_payload, strip_markdown_bold


def export_resume_to_docx(data: dict) -> io.BytesIO:
    """
    根据前端传来的/任务存储中的 RenderReadyResume 数据结构，生成格式化的 Word 文档。

    Args:
        data: RenderReadyResume 的 dict

    Returns:
        BytesIO 对象，包含文档二进制流
    """
    data = sanitize_resume_payload(data)
    doc = Document()

    # 页面边界设置
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # ==== 头部（姓名与联系方式） ====
    name = data.get("name", "姓名未定")
    head_p = doc.add_paragraph()
    head_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = head_p.add_run(name)
    run.font.size = Pt(20)
    run.font.bold = True

    contact = data.get("contact", {})
    contact_parts = []
    if contact.get("email"):
        contact_parts.append(contact["email"])
    if contact.get("phone"):
        contact_parts.append(contact["phone"])
    if contact.get("linkedin"):
        contact_parts.append(contact["linkedin"])

    if contact_parts:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_p.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = RGBColor(100, 116, 139)  # slate-500

    # 添加分割线
    doc.add_paragraph().add_run("_" * 75)

    # ==== 个人总结 ====
    if data.get("summary"):
        sec_p = doc.add_paragraph()
        run = sec_p.add_run("个人总结")
        run.font.size = Pt(12)
        run.font.bold = True

        # 简单处理 "** 加粗" 语法 (去除或加粗)
        clean_summary = strip_markdown_bold(str(data["summary"]))
        summary_p = doc.add_paragraph(clean_summary)
        summary_p.paragraph_format.space_after = Pt(12)

    # ==== 动态 Sections ====
    sections = data.get("sections", [])
    for sec in sections:
        if sec.get("type") == "summary":
            continue

        # 模块标题
        sec_p = doc.add_paragraph()
        run = sec_p.add_run(sec.get("title", "").upper())
        run.font.size = Pt(12)
        run.font.bold = True
        # 底边框用横线替代
        doc.add_paragraph().add_run("_" * 75)

        sec_type = sec.get("type")
        items = sec.get("items", [])

        if sec_type == "experience":
            for item in items:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                p.add_run(item.get("company", "")).bold = True

                # 时间居右 (通过制表符简单实现，或者直接拼接)
                p.add_run(f"    ( {item.get('duration', '')} )")

                title_p = doc.add_paragraph()
                title_p.paragraph_format.space_after = Pt(2)
                title_run = title_p.add_run(item.get("title", ""))
                title_run.font.color.rgb = RGBColor(37, 99, 235)  # blue-600
                title_run.italic = True

                for hl in item.get("highlights", []):
                    # 简单去除 Markdown 的加粗符号，在Word中纯文本展示
                    clean_hl = strip_markdown_bold(str(hl))
                    hl_p = doc.add_paragraph(clean_hl, style="List Bullet")
                    hl_p.paragraph_format.space_after = Pt(1)

        elif sec_type == "education":
            for item in items:
                p = doc.add_paragraph()
                p.add_run(item.get("school", "")).bold = True
                p.add_run(
                    f"    |    {item.get('major', '')} • {item.get('degree', '')}"
                )
                p.add_run(f"    ( {item.get('year', '')} )")

        elif sec_type == "skills":
            for item in items:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                p.add_run(item.get("category", "") + ": ").bold = True
                skills_list = item.get("skills", [])
                p.add_run(" • ".join(skills_list))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream
